from __future__ import annotations

import hashlib
import json
import logging
from pathlib import Path

from agents import Agent

from jobpipe.core.candidate_data import default_candidate_id, load_candidate_resume_json
from jobpipe.core.io import now_iso
from jobpipe.core.profile_pack import parse_profile_pack
from jobpipe.decision import (
    CandidateEvidenceContext,
    CandidateNarrativeContext,
    DecisionContext,
    build_candidate_evidence_context,
    build_candidate_narrative_context,
    build_decision_context,
    persist_candidate_materials,
)
from jobpipe.runtime.paths import primary_db_path
from jobpipe.core.primary_db import connect_primary_db, ensure_candidate, insert_generated_document
from jobpipe.model.schema import JobContext, ApplicationPackOut
from jobpipe.stages._common import run_agent, format_deadline

logger = logging.getLogger(__name__)

_PRIMARY_DB_PATH = primary_db_path()
_DEFAULT_CANDIDATE_ID = default_candidate_id()

PACK_INSTRUCTIONS = """Du er en norsk søknadsassistent. Du mottar kontekst som JSON og
produserer en komplett søknadspakke for kandidaten.

Du har tilgang til:
- job_header: stillingstittel og arbeidsgiver
- job_parsed: strukturerte jobbkrav (ansvar, must-have, nice-to-have)
- profile_match: fit_score, overlaps, gaps, hard_blockers
- pivot: pivot_score og pivot-vurdering
- moderator: endelig beslutning og cv_focus-anbefalinger
- profile_pack: kandidatens profil og mål (narrativ)
- resume_work: kandidatens arbeidshistorie med highlights (JSON Resume-format)
- resume_projects: kandidatens prosjektportefølje
- candidate_evidence_units: strukturerte, kandidatgodkjente evidensenheter
- selected_evidence_units: evidensenheter valgt for denne rollen
- decision_table: eksplisitt vurdering av can_do / can_get / should_want / can_explain
- narrative_profile: kandidatens strukturerte profesjonelle retning og kjernefortelling
- narrative_fragments: godkjente fortellingsfragmenter for CV/søknad
- job_narrative_assessment: vurdering av retning, motivasjon og pivot-troverdighet
- motivation_brief: kort, troverdig begrunnelse for hvorfor rollen gir mening nå
- narrative_strategy: deterministisk posisjoneringsanalyse (positioning_angle, brand_frame,
  why_me_now, top_value_props, cv_focus_order, cover_letter_strategy)
- advantage_assessment: fordelsanalyse (advantage_type, recruiter_hook,
  applicant_pool_hypothesis, differentiation_signals)

Bruk narrative_strategy og advantage_assessment som det primære kompasset for tonen og
vinkelen i søknadspakken:
- positioning_angle setter overordnet ramme — bruk den direkte i cover_letter_angle
- brand_frame er kandidatens én-linje identitet — la den skinne gjennom i overskrift og åpning
- why_me_now er kjernen i første avsnitt i søknadsbrevet
- recruiter_hook er åpningssetningen som stopper scanningen
- top_value_props og cv_focus_order prioriterer cv_highlights-utvalget
- cover_letter_strategy er avslutningsstrategien

Velg formulering som ligger tett på kandidatens canonical_text og respekter rewrite_policy.

For cv_highlights: velg 4-6 bullets primært fra selected_evidence_units og bruk resume_work / resume_projects
kun som støtte hvis det trengs.

For cv_highlights: velg 4-6 bullets fra selected_evidence_units og/eller resume_work.highlights / resume_projects
som matcher jobbkravene best. Omformuler lett for å speile stillingens terminologi —
men IKKE oppfinn erfaring. Hvert punkt skal stå alene og si noe konkret.

cv_highlights og cv_experience_refs MÅ ha nøyaktig samme antall elementer.

Vær troverdig og kompakt. Skriv handlingsorientert norsk.
"""


def _load_resume_context() -> dict:
    """Load JSON Resume and extract work + projects for the prompt."""
    try:
        data = load_candidate_resume_json(candidate_id=_DEFAULT_CANDIDATE_ID)
        if not data:
            return {"resume_work": [], "resume_projects": [], "resume_education": []}
        # Compact work entries: keep name, position, dates, summary, highlights
        work = []
        for w in data.get("work", []):
            work.append({
                "company": w.get("name") or w.get("company", ""),
                "position": w.get("position", ""),
                "start": w.get("startDate", ""),
                "end": w.get("endDate", "present"),
                "summary": (w.get("summary") or "")[:200],
                "highlights": w.get("highlights", []),
            })
        # Compact projects: keep name + description
        projects = [
            {"name": p.get("name", ""), "description": (p.get("description") or "")[:200]}
            for p in data.get("projects", [])
        ]
        education = [
            {
                "institution": e.get("institution", ""),
                "area": e.get("area", ""),
                "studyType": e.get("studyType", ""),
            }
            for e in data.get("education", [])
        ]
        return {"resume_work": work, "resume_projects": projects, "resume_education": education}
    except Exception as exc:  # noqa: BLE001
        logger.warning("[application_pack] could not load candidate resume context: %s", exc)
        return {"resume_work": [], "resume_projects": [], "resume_education": []}


def _application_pack_detail(ctx: JobContext) -> dict:
    profile_match = ctx.profile_match.model_dump() if ctx.profile_match else {}
    pivot = ctx.pivot.model_dump() if ctx.pivot else {}
    return {
        "overlaps": profile_match.get("overlaps", []),
        "gaps": profile_match.get("gaps", []),
        "hard_blockers": profile_match.get("hard_blockers", []),
        "match_notes": profile_match.get("notes", ""),
        "pivot_type": pivot.get("pivot_type", ""),
        "pivot_why": pivot.get("why_it_matters", []),
    }


def _application_pack_job_view(ctx: JobContext) -> dict:
    moderator = ctx.moderator.model_dump() if ctx.moderator else {}
    return {
        "title": ctx.job.get("title"),
        "employer": ctx.job.get("employer_name") or ctx.job.get("company") or "",
        "sector": ctx.job.get("sector") or "",
        "work_city": ctx.job.get("work_city") or ctx.job.get("municipal") or "",
        "work_county": ctx.job.get("work_county") or ctx.job.get("county") or "",
        "work_postalCode": ctx.job.get("work_postalCode") or "",
        "description_snip": ctx.job.get("description_snip") or ctx.job.get("description") or "",
        "triage_explanation": ctx.triage.explanation if ctx.triage else "",
        "triage_signals": ctx.triage.signals if ctx.triage else [],
        "fit_score": ctx.profile_match.fit_score if ctx.profile_match else 0,
        "pivot_score": ctx.pivot.pivot_score if ctx.pivot else 0,
        "final_decision": moderator.get("final_decision", ""),
        "recommendation_reason": moderator.get("recommendation_reason", ""),
        "detail": _application_pack_detail(ctx),
    }


def _application_pack_focus_terms(ctx: JobContext) -> list[str]:
    profile_match = ctx.profile_match.model_dump() if ctx.profile_match else {}
    moderator = ctx.moderator.model_dump() if ctx.moderator else {}
    terms: list[str] = []
    for value in moderator.get("cv_focus", []) or []:
        text = str(value).strip()
        if text:
            terms.append(text)
    for value in profile_match.get("overlaps", []) or []:
        text = str(value).strip()
        if text:
            terms.append(text)
    return terms


def _build_application_pack_contexts(
    ctx: JobContext,
    resume_ctx: dict,
) -> tuple[DecisionContext, CandidateEvidenceContext, CandidateNarrativeContext]:
    job_view = _application_pack_job_view(ctx)
    decision_context = build_decision_context(job_view, candidate_profile=parse_profile_pack(ctx.profile_pack))
    evidence_context = build_candidate_evidence_context(
        job_view,
        {
            "work": resume_ctx.get("resume_work", []),
            "projects": resume_ctx.get("resume_projects", []),
            "education": resume_ctx.get("resume_education", []),
        },
        candidate_id=_DEFAULT_CANDIDATE_ID,
        focus_terms=_application_pack_focus_terms(ctx),
        limit=6,
    )
    narrative_context = build_candidate_narrative_context(
        job_view,
        ctx.profile_pack,
        evidence_context.candidate_evidence_units,
        evidence_context.selected_evidence_units,
        candidate_id=_DEFAULT_CANDIDATE_ID,
        decision_table=decision_context.decision_table,
    )
    return decision_context, evidence_context, narrative_context


def _narrative_strategy_payload(ctx: JobContext) -> dict:
    """Compact narrative_strategy_v3 fields for the LLM prompt."""
    ns = ctx.narrative_strategy_v3
    if not ns:
        return {}
    return {
        "positioning_angle": ns.positioning_angle,
        "brand_frame": ns.brand_frame,
        "why_me_now": ns.why_me_now,
        "top_value_props": ns.top_value_props,
        "objections_to_handle": ns.objections_to_handle,
        "cv_focus_order": ns.cv_focus_order,
        "cover_letter_strategy": ns.cover_letter_strategy,
    }


def _advantage_assessment_payload(ctx: JobContext) -> dict:
    """Compact advantage_assessment_v3 fields for the LLM prompt."""
    aa = ctx.advantage_assessment_v3
    if not aa:
        return {}
    return {
        "advantage_type": aa.advantage_type,
        "recruiter_hook": aa.recruiter_hook,
        "applicant_pool_hypothesis": aa.applicant_pool_hypothesis,
        "differentiation_signals": aa.differentiation_signals,
        "advantage_signals": aa.advantage_signals,
        "objection_signals": aa.objection_signals,
        "stretch_level": aa.stretch_level,
    }


def _build_application_pack_payload(
    ctx: JobContext,
    resume_ctx: dict,
    *,
    decision_context: DecisionContext | None = None,
    evidence_context: CandidateEvidenceContext | None = None,
    narrative_context: CandidateNarrativeContext | None = None,
) -> dict:
    decision_context, evidence_context, narrative_context = (
        decision_context,
        evidence_context,
        narrative_context,
    )
    if decision_context is None or evidence_context is None or narrative_context is None:
        decision_context, evidence_context, narrative_context = _build_application_pack_contexts(ctx, resume_ctx)
    return {
        "job_header": {
            "title": ctx.job.get("title"),
            "employer_name": ctx.job.get("employer_name"),
            "sector": ctx.job.get("sector"),
            "deadline": format_deadline((ctx.job.get("applicationDue") or "").strip()),
            "source_url": ctx.job.get("sourceurl") or ctx.job.get("link"),
        },
        "job_parsed": ctx.parsed.model_dump() if ctx.parsed else {},
        "profile_match": ctx.profile_match.model_dump() if ctx.profile_match else {},
        "pivot": ctx.pivot.model_dump() if ctx.pivot else {},
        "moderator": ctx.moderator.model_dump() if ctx.moderator else {},
        "narrative_strategy": _narrative_strategy_payload(ctx),
        "advantage_assessment": _advantage_assessment_payload(ctx),
        "profile_pack": ctx.profile_pack[:3000],
        "resume_work": resume_ctx["resume_work"],
        "resume_projects": resume_ctx["resume_projects"],
        "resume_education": resume_ctx.get("resume_education", []),
        "candidate_evidence_units": [unit.model_dump(mode="json") for unit in evidence_context.candidate_evidence_units],
        "selected_evidence_units": [selection.model_dump(mode="json") for selection in evidence_context.selected_evidence_units],
        "decision_table": decision_context.decision_table.model_dump(mode="json"),
        "narrative_profile": narrative_context.narrative_profile.model_dump(mode="json"),
        "narrative_fragments": [fragment.model_dump(mode="json") for fragment in narrative_context.narrative_fragments],
        "narrative_evidence_links": [link.model_dump(mode="json") for link in narrative_context.narrative_evidence_links],
        "job_narrative_assessment": narrative_context.job_narrative_assessment.model_dump(mode="json"),
        "motivation_brief": narrative_context.job_narrative_assessment.motivation_brief,
    }


def _preview_text(pack_data: dict) -> str:
    parts: list[str] = []
    headline = (pack_data.get("positioning_headline") or "").strip()
    if headline:
        parts.append(headline)
    cover_angle = (pack_data.get("cover_letter_angle") or "").strip()
    if cover_angle:
        parts.append(cover_angle)
    highlights = [str(x).strip() for x in (pack_data.get("cv_highlights") or []) if str(x).strip()]
    if highlights:
        parts.append(" | ".join(highlights[:3]))
    return " ".join(parts)[:800]


def _document_id(candidate_id: str, job_id: str, kind: str, storage_path: Path) -> str:
    raw = f"{candidate_id}|{job_id}|{kind}|{storage_path.name}"
    return "doc_" + hashlib.sha1(raw.encode("utf-8")).hexdigest()[:20]


def _generated_document_rows(
    ctx: JobContext,
    pack_data: dict,
    draft_path: Path,
    docx_path: Path | None,
    *,
    candidate_id: str,
    generated_at: str,
) -> list[dict]:
    evaluation_id = f"{ctx.meta.run_id}:{ctx.job_id}"
    preview = _preview_text(pack_data)
    rows = [
        {
            "document_id": _document_id(candidate_id, ctx.job_id, "application_pack_json", draft_path),
            "candidate_id": candidate_id,
            "job_id": ctx.job_id,
            "evaluation_id": evaluation_id,
            "kind": "application_pack_json",
            "producer": "jobpipe_pipeline",
            "status": "draft",
            "storage_path": str(draft_path.resolve()),
            "preview_text": preview,
            "document_json": pack_data,
            "created_at": generated_at,
            "updated_at": generated_at,
        }
    ]
    if docx_path and docx_path.exists():
        rows.append(
            {
                "document_id": _document_id(candidate_id, ctx.job_id, "cv_highlights_docx", docx_path),
                "candidate_id": candidate_id,
                "job_id": ctx.job_id,
                "evaluation_id": evaluation_id,
                "kind": "cv_highlights_docx",
                "producer": "jobpipe_pipeline",
                "status": "draft",
                "storage_path": str(docx_path.resolve()),
                "preview_text": preview,
                "document_json": {
                    "positioning_headline": pack_data.get("positioning_headline", ""),
                    "cv_highlights": pack_data.get("cv_highlights", []),
                    "cv_experience_refs": pack_data.get("cv_experience_refs", []),
                },
                "created_at": generated_at,
                "updated_at": generated_at,
            }
        )
    return rows


def _register_generated_documents(conn, rows: list[dict]) -> None:
    for row in rows:
        insert_generated_document(conn, row)


def _persist_candidate_materials_for_pack(
    conn,
    *,
    ctx: JobContext,
    evidence_context: CandidateEvidenceContext | None,
    narrative_context: CandidateNarrativeContext | None,
    updated_at: str,
    candidate_id: str,
) -> None:
    if evidence_context is None or narrative_context is None:
        return

    persist_candidate_materials(
        conn,
        candidate_id=candidate_id,
        job_id=ctx.job_id,
        evaluation_id=f"{ctx.meta.run_id}:{ctx.job_id}",
        evidence_context=evidence_context,
        narrative_context=narrative_context,
        updated_at=updated_at,
    )


def _sync_generated_documents(
    ctx: JobContext,
    pack_data: dict,
    draft_path: Path,
    docx_path: Path | None = None,
    *,
    decision_context: DecisionContext | None = None,
    evidence_context: CandidateEvidenceContext | None = None,
    narrative_context: CandidateNarrativeContext | None = None,
) -> None:
    try:
        conn = connect_primary_db(_PRIMARY_DB_PATH)
        try:
            ensure_candidate(conn, candidate_id=_DEFAULT_CANDIDATE_ID)
            now = now_iso()
            _register_generated_documents(
                conn,
                _generated_document_rows(
                    ctx,
                    pack_data,
                    draft_path,
                    docx_path,
                    candidate_id=_DEFAULT_CANDIDATE_ID,
                    generated_at=now,
                ),
            )
            _persist_candidate_materials_for_pack(
                conn,
                ctx=ctx,
                evidence_context=evidence_context,
                narrative_context=narrative_context,
                updated_at=now,
                candidate_id=_DEFAULT_CANDIDATE_ID,
            )

            conn.commit()
        finally:
            conn.close()
    except Exception as exc:  # noqa: BLE001
        logger.warning("[application_pack] could not sync generated_documents metadata: %s", exc)


def _generate_cv_docx(pack_data: dict, job_input: dict, job_path: Path) -> Path | None:
    """Generate a DOCX 'Relevant Experience' supplement from cv_highlights."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("[application_pack] python-docx not installed — skipping DOCX generation")
        return None

    highlights = pack_data.get("cv_highlights", [])
    refs = pack_data.get("cv_experience_refs", [])
    if not highlights:
        return None

    title = (job_input.get("title") or "").strip()
    employer = (job_input.get("employer_name") or job_input.get("company") or "").strip()
    headline = pack_data.get("positioning_headline", "")
    cover_angle = pack_data.get("cover_letter_angle", "")
    interview_prep = pack_data.get("interview_prep", [])

    doc = Document()

    # A4, ~2 cm margins
    for section in doc.sections:
        section.page_width = 11906
        section.page_height = 16838
        margin = int(0.79 * 1440)
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    # Header: name + role
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Lars Værland")
    run.bold = True
    run.font.size = Pt(16)

    if title or employer:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = f"{title}" + (f" — {employer}" if employer else "")
        r2 = p2.add_run(label)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    if headline:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(headline)
        r3.italic = True
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Relevant experience
    h1 = doc.add_paragraph()
    rh = h1.add_run("Relevante erfaringspunkter")
    rh.bold = True
    rh.font.size = Pt(13)

    for i, bullet in enumerate(highlights):
        ref = refs[i] if i < len(refs) else ""
        p_b = doc.add_paragraph(style="List Bullet")
        p_b.add_run(bullet)
        if ref:
            p_b.add_run(f"  [{ref}]").font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Cover letter angle
    if cover_angle:
        h2 = doc.add_paragraph()
        rh2 = h2.add_run("Søknadsvinkel")
        rh2.bold = True
        rh2.font.size = Pt(13)
        doc.add_paragraph(cover_angle)
        doc.add_paragraph()

    # Interview prep
    if interview_prep:
        h3 = doc.add_paragraph()
        rh3 = h3.add_run("Intervjuforberedelse")
        rh3.bold = True
        rh3.font.size = Pt(13)
        for q in interview_prep:
            p_q = doc.add_paragraph(style="List Bullet")
            p_q.add_run(q)

    out_path = job_path / "07_cv_highlights.docx"
    doc.save(str(out_path))
    logger.info("[application_pack] saved cv_highlights DOCX to %s", out_path)
    return out_path


def application_pack_stage_factory(model: str, web_search: bool = False):  # noqa: ARG001

    agent = Agent(
        name="application_pack_agent",
        model=model,
        instructions=PACK_INSTRUCTIONS,
        output_type=ApplicationPackOut,
    )
    resume_ctx = _load_resume_context()

    def should_run(ctx: JobContext) -> bool:
        return bool(ctx.moderator and ctx.moderator.final_decision in ("APPLY_STRONGLY", "APPLY"))

    def run(ctx: JobContext, job_dir: str) -> JobContext:
        job_path = Path(job_dir)

        decision_context, evidence_context, narrative_context = _build_application_pack_contexts(ctx, resume_ctx)
        payload = _build_application_pack_payload(
            ctx,
            resume_ctx,
            decision_context=decision_context,
            evidence_context=evidence_context,
            narrative_context=narrative_context,
        )

        input_text = "Kontekst (JSON):\n" + json.dumps(payload, ensure_ascii=False, indent=2)
        logger.info("[application_pack] running for job %s", ctx.job_id)

        result = run_agent(agent, input_text, trace={"stage": "application_pack", "job_id": ctx.job_id})
        pack = result.final_output_as(ApplicationPackOut)
        ctx.application_pack = pack

        # Persist the draft for inspection
        pack_data = pack.model_dump()
        draft_path = job_path / "application_pack_draft.json"
        draft_path.write_text(json.dumps(pack_data, ensure_ascii=False, indent=2), encoding="utf-8")

        logger.info(
            "[application_pack] done for job %s (%d cv_highlights)",
            ctx.job_id, len(pack_data.get("cv_highlights", [])),
        )

        # Generate DOCX supplement
        docx_path = _generate_cv_docx(pack_data, ctx.job, job_path)
        _sync_generated_documents(
            ctx,
            pack_data,
            draft_path=draft_path,
            docx_path=docx_path,
            decision_context=decision_context,
            evidence_context=evidence_context,
            narrative_context=narrative_context,
        )

        return ctx

    return should_run, run
